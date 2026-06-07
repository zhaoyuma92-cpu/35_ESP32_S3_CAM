from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ── 页面边距 ──────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── 辅助函数 ─────────────────────────────────────────────
def heading(text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_before = Pt(14 if level == 1 else 8)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold_parts=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        p.add_run(text)
    return p

def code_block(text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1)
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(text)
    run.font.name  = 'Courier New'
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x20, 0x60, 0x20)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  'F0F4F0')
    p._p.pPr.append(shading)
    return p

def bullet(text, bold_parts=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    if bold_parts:
        parts = text.split('**')
        for i, part in enumerate(parts):
            run = p.add_run(part)
            run.bold = (i % 2 == 1)
    else:
        p.add_run(text)
    return p

def tip_box(label, text, bg='FFF8E1', label_color=(0xE6, 0x5C, 0x00)):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.8)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    r1 = p.add_run(f'【{label}】 ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(*label_color)
    r2 = p.add_run(text)
    r2.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    shading = OxmlElement('w:shd')
    shading.set(qn('w:val'),   'clear')
    shading.set(qn('w:color'), 'auto')
    shading.set(qn('w:fill'),  bg)
    p._p.pPr.append(shading)
    return p

def divider():
    p = doc.add_paragraph('─' * 55)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xCC, 0xCC, 0xCC)
        run.font.size = Pt(9)

# ═══════════════════════════════════════════════════════════
# 封面
# ═══════════════════════════════════════════════════════════
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_p.paragraph_format.space_before = Pt(30)
tr = title_p.add_run('ESP32-P4-NANO 固件调试报告')
tr.bold = True
tr.font.size = Pt(22)
tr.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)

sub_p = doc.add_paragraph()
sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub_p.add_run('烧录遇到的问题、原因与修复过程（初学者版）')
sr.font.size = Pt(13)
sr.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_p.add_run(f'生成日期：{datetime.date.today()}').font.size = Pt(10)

doc.add_page_break()

# ═══════════════════════════════════════════════════════════
# 第 0 节：读者须知
# ═══════════════════════════════════════════════════════════
heading('写在前面：这份报告是干什么的', 1, (0x1A, 0x53, 0x76))
body('你刚刚完成了一次嵌入式固件的"编译 → 烧录 → 调试"全流程。过程中遇到了好几个错误——这很正常，每一个错误背后都有一个可以学到东西的原因。\n\n这份报告会用最通俗的语言，把每个问题解释清楚，让你下次再遇到类似情况时有底气。')

tip_box('读法建议',
        '建议从第 1 节顺序读下来。每一节都有"是什么错 → 为什么出错 → 怎么修"三步结构。遇到不懂的术语，看括号里的解释就够了。',
        'E8F5E9', (0x2E, 0x7D, 0x32))

divider()

# ═══════════════════════════════════════════════════════════
# 第 1 节：项目是什么
# ═══════════════════════════════════════════════════════════
heading('第 1 节  这个项目是做什么的', 1, (0x1A, 0x53, 0x76))

body('先了解你在做什么，才能理解为什么会出错。')

heading('硬件', 2)
bullet('**开发板**：ESP32-P4-NANO（芯片：ESP32-P4，双核 RISC-V，主频 360 MHz）')
bullet('**摄像头接口**：MIPI-CSI（连接 OV5647 图像传感器，暂未实现真实驱动）')
bullet('**存储**：16 MB Flash + 32 MB PSRAM（板载外部高速内存）')
bullet('**串口**：CH343 USB 转串口芯片，对应电脑上的 COM8')

heading('软件目标', 2)
body('固件实现一套**视觉位移测量系统**：')
bullet('以 30 帧/秒连续采集摄像头图像（现阶段用"假相机"生成测试图片）')
bullet('在每张图片里追踪最多 4 个目标点的位置，计算它们的位移量（精度达到 1/256 像素）')
bullet('把每帧的结果写入 SD 卡上的 CSV 文件，便于后续分析')

heading('固件整体流程（状态机）', 2)
body('固件启动后按以下顺序运行，每个阶段在串口里都有日志输出：')
code_block(
    'BOOT → CAMERA_INIT → IDLE → RECORDING → FLUSHING → FINISHED\n'
    '启动     初始化摄像头    空闲    录制中         写入完成       正常结束'
)

divider()

# ═══════════════════════════════════════════════════════════
# 第 2 节：问题 1 —— 串口占用
# ═══════════════════════════════════════════════════════════
heading('第 2 节  问题一：串口被占用，烧录失败', 1, (0xC6, 0x28, 0x28))

heading('报错内容', 2)
code_block('A fatal error occurred: Could not open COM8, the port is busy or doesn\'t exist.\n(could not open port \'COM8\': PermissionError(13, \'拒绝访问。\', None, 5))')

heading('这是什么意思？', 2)
body('COM8 是电脑和 ESP32-P4 之间的"通信管道"。烧录工具（esptool）需要独占这个管道才能写入程序。\n\n就好比你要往一个水管里灌水，但别人的水龙头正开着——水管被占着，你灌不进去。')

heading('原因', 2)
body('之前在 VS Code 里打开过**串口监视器**（idf.py monitor），关闭窗口时进程没有完全退出，后台残留了 **3 个 idf_monitor 进程**，全都占着 COM8：')
code_block(
    'PID 1904  : esp_idf_monitor -p COM8 -b 115200 ...\n'
    'PID 16244 : idf_monitor.py  -p COM8 -b 115200 ...\n'
    'PID 16272 : esp_idf_monitor -p COM8 -b 115200 ...'
)

heading('修复方法', 2)
body('直接关掉这 3 个进程，COM8 就空出来了。之后烧录立刻成功。')
tip_box('以后怎么避免',
        '每次烧录前，先在 VS Code 里按 Ctrl+C 关掉串口监视器，或者在终端里用 Ctrl+] 退出 idf.py monitor，确保 COM8 没有其他程序在用。',
        'FFF3E0', (0xE6, 0x5C, 0x00))

divider()

# ═══════════════════════════════════════════════════════════
# 第 3 节：问题 2 —— 内存不足
# ═══════════════════════════════════════════════════════════
heading('第 3 节  问题二：内存不足，相机初始化失败', 1, (0xC6, 0x28, 0x28))

heading('报错内容', 2)
code_block(
    'E p4_camera: fake frame allocation failed: 1228800 bytes\n'
    'E main: camera init failed: ESP_ERR_NO_MEM\n'
    'I node_state: state=ERROR'
)

heading('这是什么意思？', 2)
body('固件里有一个"假相机"功能——在真实摄像头驱动还没写好之前，它自己生成一张假图片来测试后续流程。')
body('每张假图片的大小 = 宽度 × 高度 = 1280 × 960 = **1,228,800 字节（约 1.2 MB）**。')
body('但问题是：这张图片需要存放在内存里，而芯片的内部内存（SRAM）只有约 550 KB，**根本装不下 1.2 MB 的图片**。')

heading('内存概念类比', 2)
tip_box('类比',
        '内部 SRAM 像桌面，只有 550KB 大小。要摆一张 1.2MB 的"图纸"根本放不下。\n'
        'PSRAM 像旁边的一个大抽屉，有 32MB，可以放很多东西——但这个抽屉当时还没打开。',
        'E3F2FD', (0x15, 0x65, 0xA3))

heading('修复方法', 2)
body('**临时方案**：在 PSRAM 还没开启时，把假相机分辨率从 1280×960 缩小到 320×240（= 76,800 字节，只需 75KB），完全放得进内部 SRAM。')
body('同时在 Kconfig（项目配置文件）里加了两个可调整的配置项，方便以后随时切换分辨率：')
code_block(
    'DISP_FAKE_FRAME_WIDTH  = 1280  (开启PSRAM后恢复)\n'
    'DISP_FAKE_FRAME_HEIGHT = 960'
)
body('还把 4 个追踪目标的 ROI（感兴趣区域）坐标按比例缩放，确保在 320×240 的图里仍然有效。')

divider()

# ═══════════════════════════════════════════════════════════
# 第 4 节：问题 3 —— 芯片版本拒绝
# ═══════════════════════════════════════════════════════════
heading('第 4 节  问题三：芯片版本太旧，bootloader 拒绝运行', 1, (0xC6, 0x28, 0x28))

heading('报错内容', 2)
code_block(
    'A fatal error occurred: bootloader/bootloader.bin requires\n'
    'chip revision in range [v3.1 - v3.99]\n'
    '(this chip is revision v1.3). Use --force to flash anyway.'
)

heading('这是什么意思？', 2)
body('ESP32-P4 这款芯片有两个不兼容的"硅片版本"家族：')
bullet('**v0.x / v1.x**：早期工程样片，用于研发测试，硬件电路和寄存器布局较老')
bullet('**v3.x**：量产正式版，硬件设计做了大改，与 v1.x 不兼容')
body('你手上这块板子的芯片是 **v1.3（早期样片）**。但 IDF v5.5.4 默认编译出来的 bootloader 只支持 v3.x 芯片，遇到 v1.3 直接拒绝启动。')

tip_box('什么是 bootloader？',
        'Bootloader（引导程序）是芯片上电后最先运行的一小段代码，负责检查硬件、加载你写的固件。'
        '它就像电脑开机时的 BIOS。如果它不认识你的芯片版本，固件根本无法运行。',
        'F3E5F5', (0x6A, 0x1B, 0x9A))

heading('另一个触发因素：Flash 频率', 2)
body('我们同时把 Flash 大小从默认 2MB 改为实际的 16MB，这个改动自动把 Flash 读取频率从 40MHz 升级到 80MHz。'
     '而 80MHz Flash 频率也需要 v3.x 芯片，形成双重拒绝。')

heading('修复方法', 2)
body('在 sdkconfig.defaults（项目默认配置文件）里显式告诉 IDF：')
bullet('**使用 pre-v3 支持路径**（`ESP32P4_SELECTS_REV_LESS_V3=y`）')
bullet('**最低支持版本设为 v1.0**（`ESP32P4_REV_MIN_100=y`）')
bullet('**Flash 频率回到 40MHz**（`ESPTOOLPY_FLASHFREQ_40M=y`）')
body('修改后的 sdkconfig.defaults 最终内容：')
code_block(
    'CONFIG_IDF_TARGET="esp32p4"\n'
    'CONFIG_FREERTOS_HZ=1000\n'
    'CONFIG_COMPILER_OPTIMIZATION_PERF=y\n\n'
    '# Flash: 实际 16MB，保持 40MHz 与 v1.x 兼容\n'
    'CONFIG_ESPTOOLPY_FLASHSIZE_16MB=y\n'
    'CONFIG_ESPTOOLPY_FLASHFREQ_40M=y\n\n'
    '# 芯片版本：这块板子是 v1.3（pre-v3 早期样片）\n'
    'CONFIG_ESP32P4_SELECTS_REV_LESS_V3=y\n'
    'CONFIG_ESP32P4_REV_MIN_100=y\n\n'
    '# PSRAM：HEX 模式 80MHz\n'
    'CONFIG_SPIRAM=y\n'
    'CONFIG_SPIRAM_MODE_HEX=y\n'
    'CONFIG_SPIRAM_SPEED_80M=y'
)

divider()

# ═══════════════════════════════════════════════════════════
# 第 5 节：PSRAM 开启
# ═══════════════════════════════════════════════════════════
heading('第 5 节  PSRAM 启用成功', 1, (0x1B, 0x5E, 0x20))

heading('PSRAM 是什么？', 2)
body('PSRAM（伪静态随机存储器）是焊在 ESP32-P4-NANO 板子上的**外部扩展内存**，容量 32MB。')
body('它通过一条高速总线（HEX/16线模式，最高 200MHz）连接到芯片，读写速度比普通 SPI 快得多，但比芯片内部 SRAM 慢一些。')

tip_box('类比',
        '内部 SRAM ≈ CPU 旁边的寄存器纸（极快，只有 550KB）\n'
        'PSRAM ≈ 旁边的大笔记本（稍慢，但有 32MB）\n'
        'Flash ≈ 硬盘（最慢，16MB，断电不丢失）',
        'E8F5E9', (0x2E, 0x7D, 0x32))

heading('为什么之前没开启？', 2)
body('sdkconfig.defaults 里没有 CONFIG_SPIRAM=y，IDF 默认不初始化 PSRAM。芯片上电后 PSRAM 硬件是存在的，但软件不去初始化它，就当它不存在。')

heading('开启后的启动日志', 2)
code_block(
    'I hex_psram: vendor id  : 0x0d (AP)         ← AP Memory 品牌\n'
    'I hex_psram: density    : 0x07 (256 Mbit)   ← 256Mb = 32MB\n'
    'I hex_psram: good-die   : 0x06 (Pass)       ← 良品测试通过\n'
    'I esp_psram: Found 32MB PSRAM device\n'
    'I esp_psram: Speed: 80MHz\n'
    'I esp_psram: SPI SRAM memory test OK        ← 内存自检通过\n'
    'I esp_psram: Adding pool of 32768K of PSRAM memory to heap allocator\n'
    '                                             ← 32MB 加入可用内存池'
)
body('开启后，heap_caps_malloc(size, MALLOC_CAP_SPIRAM) 就能从这 32MB 里分配内存，1280×960 的帧缓冲（1.2MB）轻松放下。')

divider()

# ═══════════════════════════════════════════════════════════
# 第 6 节：WDT 警告
# ═══════════════════════════════════════════════════════════
heading('第 6 节  运行时警告：任务看门狗（WDT）', 1, (0xE6, 0x5C, 0x00))

heading('报错内容', 2)
code_block(
    'E task_wdt: Task watchdog got triggered.\n'
    'E task_wdt:  - IDLE0 (CPU 0)\n'
    'E task_wdt: Tasks currently running:\n'
    'E task_wdt: CPU 0: CameraTask'
)

heading('这是什么意思？', 2)
body('FreeRTOS（固件用的实时操作系统）有一个"看门狗"机制：它要求每隔一段时间，空闲任务（IDLE）必须得到 CPU 运行的机会。'
     '如果某个任务长时间占着 CPU 不让步，看门狗就会报警。')
body('这里是 CameraTask（摄像头采集任务）在 CPU0 上运行了太长时间，导致 IDLE 任务被饿死。')

heading('根本原因：PSRAM 访问速度', 2)
body('1280×960 的帧缓冲存在 PSRAM 里。ROI 追踪算法每帧需要读取最多 4 个 150×150 像素区域，逐像素遍历——这涉及大量随机访问 PSRAM。')
body('**PSRAM 访问延迟**比内部 SRAM 高约 10-20 倍，导致单帧处理时间从理想的 33ms 拉长到约 46ms：')
code_block(
    'dt_min = 44.9ms\n'
    'dt_avg = 46.3ms  （目标 33.3ms，慢了 40%）\n'
    'dt_max = 182ms   （偶发的极端情况）'
)

heading('严重吗？', 2)
body('**不严重**。300 帧全部正常处理完，状态机最终到达 FINISHED，系统没有重启。WDT 默认只打印警告，不会崩溃。')
tip_box('生产环境如何解决',
        '把帧缓冲放在 PSRAM，但 ROI 处理时先用 DMA 把感兴趣的小区域（~150×150 像素）搬到内部 SRAM，再做像素运算。'
        '这样可以兼顾"大容量缓冲（PSRAM）"和"高速运算（SRAM）"两个优势。',
        'FFF3E0', (0xE6, 0x5C, 0x00))

divider()

# ═══════════════════════════════════════════════════════════
# 第 7 节：最终成功日志
# ═══════════════════════════════════════════════════════════
heading('第 7 节  最终成功运行的串口日志', 1, (0x1B, 0x5E, 0x20))
body('修复所有问题后，完整的运行日志如下（精简版）：')
code_block(
    'I boot: chip revision: v1.3                   ← v1.3 被接受，不再拒绝\n'
    'I boot.esp32p4: SPI Speed  : 40MHz\n'
    'I boot.esp32p4: SPI Flash Size : 16MB         ← 正确识别 16MB Flash\n'
    '\n'
    'I esp_psram: Found 32MB PSRAM device\n'
    'I esp_psram: Speed: 80MHz\n'
    'I esp_psram: SPI SRAM memory test OK          ← PSRAM 自检通过\n'
    'I esp_psram: Adding pool of 32768K ...\n'
    '\n'
    'I efuse_init: Min chip rev: v1.0\n'
    'I efuse_init: Chip rev:     v1.3              ← 满足最低版本要求\n'
    '\n'
    'I node_state: state=BOOT\n'
    'I main: ESP32-P4-NANO firmware=v0.1.0\n'
    'I node_state: state=CAMERA_INIT\n'
    'W p4_camera: using fake camera backend         ← 假相机模式（正常）\n'
    'I p4_camera: fake frame: 1280x960 @ 30fps     ← 全分辨率，从PSRAM分配\n'
    'I node_state: state=RECORDING\n'
    'W csv: open failed: /sdcard/displacement.csv  ← 无SD卡，正常跳过\n'
    '\n'
    'I write_task: done frames=300 batches=10      ← 10秒300帧全部完成\n'
    'I node_state: state=FINISHED                  ← 正常结束 ✓'
)

divider()

# ═══════════════════════════════════════════════════════════
# 第 8 节：所有修改的文件汇总
# ═══════════════════════════════════════════════════════════
heading('第 8 节  本次修改了哪些文件', 1, (0x1A, 0x53, 0x76))

rows = [
    ('文件', '修改内容', '原因'),
    ('sdkconfig.defaults',
     '加入 Flash 16MB、Flash 40MHz、\nSELECTS_REV_LESS_V3、REV_MIN_100、\nSPIRAM/SPIRAM_SPEED_80M',
     '修复芯片版本拒绝 + 开启 PSRAM'),
    ('main/Kconfig.projbuild',
     '新增 DISP_FAKE_FRAME_WIDTH/HEIGHT\n配置项，默认 1280×960',
     '让假相机分辨率可配置'),
    ('main/drivers/p4_camera.c',
     '假相机模式下使用 Kconfig 分辨率；\n优先从 PSRAM 分配帧缓冲',
     '支持 PSRAM 大帧缓冲'),
    ('main/config/app_config.c',
     '假相机模式下 ROI 坐标按\n配置分辨率自动缩放',
     '避免 ROI 超出图像边界'),
]

table = doc.add_table(rows=len(rows), cols=3)
table.style = 'Table Grid'
widths = [Cm(4.5), Cm(7.5), Cm(5.5)]
for i, row_data in enumerate(rows):
    row = table.rows[i]
    for j, cell_text in enumerate(row_data):
        cell = row.cells[j]
        cell.width = widths[j]
        p = cell.paragraphs[0]
        run = p.add_run(cell_text)
        run.font.size = Pt(9.5)
        if i == 0:
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            shading = OxmlElement('w:shd')
            shading.set(qn('w:val'),   'clear')
            shading.set(qn('w:color'), 'auto')
            shading.set(qn('w:fill'),  '1A5376')
            cell._tc.get_or_add_tcPr().append(shading)

divider()

# ═══════════════════════════════════════════════════════════
# 第 9 节：关键知识点速查
# ═══════════════════════════════════════════════════════════
heading('第 9 节  关键知识点速查', 1, (0x1A, 0x53, 0x76))

body('以下是本次出现的技术术语，以后遇到可以快速回来查：')

terms = [
    ('COM 端口',       '电脑给每个 USB 串口设备分配的编号（如 COM8）。多个程序不能同时占用同一个 COM 口。'),
    ('esptool',        'Espressif 官方的烧录工具，负责把编译好的固件通过串口写入 Flash。'),
    ('bootloader',     '芯片上电最先运行的引导程序，负责硬件检查和加载主固件。类似电脑的 BIOS。'),
    ('sdkconfig',      'IDF 项目的配置数据库文件，所有 CONFIG_XXX 选项都存在里面。不要手动删除，除非要重置配置。'),
    ('sdkconfig.defaults', '你自己写的"默认值文件"，每次重新生成 sdkconfig 时从这里读取初始值。应该提交到 git。'),
    ('SRAM',           '芯片内部的高速内存（550KB），CPU 直接访问，速度极快。'),
    ('PSRAM',          '板子上外挂的扩展内存（32MB），通过总线访问，比 SRAM 慢但容量大。'),
    ('MALLOC_CAP_SPIRAM', 'esp_heap_caps_malloc 的标志，表示从 PSRAM 分配内存。'),
    ('FreeRTOS',       '固件用的实时操作系统，管理多个并发任务（Task）的调度。'),
    ('Task WDT',       '任务看门狗。确保每个任务都能定期获得 CPU 时间，防止某个任务独占 CPU 导致系统假死。'),
    ('chip revision',  'ESP32-P4 的硅片版本号。v1.x 是工程样片，v3.x 是量产版，两者硬件不兼容。'),
    ('ROI',            'Region of Interest，感兴趣区域。在图像处理中指定只分析图像的某个矩形区域。'),
    ('Q8 定点数',       '一种不用浮点运算表示小数的方法：把真实坐标乘以 256，存成整数。例如 cx_q8=40960 表示 x=160.0 像素。'),
    ('idf.py monitor', '打开串口监视器查看固件日志的命令。使用完要记得用 Ctrl+] 退出，否则占用 COM 口。'),
]

for term, desc in terms:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r1 = p.add_run(f'{term}：')
    r1.bold = True
    r1.font.color.rgb = RGBColor(0x1A, 0x53, 0x76)
    p.add_run(desc)

divider()

# ═══════════════════════════════════════════════════════════
# 第 10 节：下一步
# ═══════════════════════════════════════════════════════════
heading('第 10 节  后续可以做什么', 1, (0x1A, 0x53, 0x76))

body('固件现在处于"能跑的最小可用状态"，以下是自然的下一步方向：')

heading('短期（1-2 周）', 2)
bullet('**插 SD 卡**：把格式化为 FAT32 的 SD 卡插上，验证 CSV 文件能否正常写入')
bullet('**观察 CSV 数据**：用 Excel 打开，查看 4 个目标的 dx/dy 位移数据是否呈现预期的三角波形')
bullet('**调整帧率和时长**：在 app_config.c 里修改 frame_rate_hz 和 duration_s，看看不同配置下的行为')

heading('中期（1-2 月）', 2)
bullet('**实现真实摄像头驱动**：p4_camera.c 里有一个 TODO，需要写 MIPI-CSI + OV5647 的初始化代码')
bullet('**DMA 优化 ROI 处理**：把 ROI 区域从 PSRAM 拷贝到 SRAM 再处理，消除 WDT 警告，达到 30fps')
bullet('**HTTP 远程控制**：main.c 里有一处 TODO 注释，预留了 HTTP 接口用于远程触发采集')

heading('工具建议', 2)
bullet('用 **idf.py menuconfig** 可视化地浏览所有配置项，不用手动编辑 sdkconfig.defaults')
bullet('用 **idf.py monitor** 看实时日志，按 Ctrl+] 退出（不是 Ctrl+C）')
bullet('用 **idf.py size** 查看固件各部分占用了多少内存')

# ── 尾部 ─────────────────────────────────────────────────
doc.add_paragraph()
end_p = doc.add_paragraph()
end_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
er = end_p.add_run('── 报告结束 ──')
er.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
er.font.size = Pt(10)

out = r'd:\Users\zhaoy\Desktop\python_c_learning\MCU\35_ESP32_S3_CAM\ESP32-P4-NANO调试报告.docx'
doc.save(out)
print(f'saved: {out}')
